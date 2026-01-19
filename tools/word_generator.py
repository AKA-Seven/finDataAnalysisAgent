# tools/word_generator.py
"""
Word报告生成工具（支持模板补全、表格插入）
输入JSON指定sections（heading/paragraph/table）
"""

import json
from docx import Document
from docx.shared import Inches
from utils.file_utils import get_output_path, df_to_table_data
from tools.base_tool import BaseTool

class WordGeneratorTool(BaseTool):
    name = "word_generator"
    description = """生成或补全Word报告。
输入JSON: {
  "template_path": "可选模板路径（如templates/report.docx）",
  "sections": [{"type": "heading"/"paragraph"/"table", "text"/"data"/"caption": ...}],
  "output_filename": "report.docx"
}"""

    def run(self, input_str: str) -> str:
        try:
            data = json.loads(input_str)
            template_path = data.get("template_path")
            sections = data.get("sections", [])
            output_filename = data.get("output_filename", "report.docx")
        except json.JSONDecodeError:
            return "错误：输入必须为有效JSON"

        try:
            if template_path:
                doc = Document(template_path)
            else:
                doc = Document()

            for sec in sections:
                sec_type = sec.get("type")
                if sec_type == "heading":
                    doc.add_heading(sec.get("text", ""), level=sec.get("level", 1))
                elif sec_type == "paragraph":
                    doc.add_paragraph(sec.get("text", ""))
                elif sec_type == "table":
                    table_data = sec.get("data", [])
                    caption = sec.get("caption", "")
                    if table_data:
                        rows, cols = len(table_data), len(table_data[0])
                        table = doc.add_table(rows=rows, cols=cols)
                        for i, row_data in enumerate(table_data):
                            row_cells = table.rows[i].cells
                            for j, cell in enumerate(row_data):
                                row_cells[j].text = str(cell)
                        if caption:
                            doc.add_paragraph(caption)

            output_path = get_output_path(output_filename)
            doc.save(str(output_path))
            return f"Word报告生成成功！文件路径：{output_path}"
        except Exception as e:
            return f"Word生成错误：{str(e)}"