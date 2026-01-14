from docx import Document
from docx.shared import Inches
from typing import Optional, List, Any

class WordOperator:
    """Word文档操作核心类，封装基础读写业务"""
    def __init__(self):
        self.document = None
        self.file_path = None

    def new_word(self) -> bool:
        """新建Word文档"""
        try:
            self.document = Document()
            return True
        except Exception as e:
            raise Exception(f"新建Word失败：{str(e)}")

    def open_word(self, file_path: str) -> bool:
        """打开已存在的Word文档"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Word文件不存在：{file_path}")
        try:
            self.document = Document(file_path)
            self.file_path = file_path
            return True
        except Exception as e:
            raise Exception(f"打开Word失败：{str(e)}")

    def add_paragraph(self, text: str, save_immediately: bool = False) -> bool:
        """添加段落文本"""
        if not self.document:
            raise Exception("请先新建或打开Word文档")
        try:
            self.document.add_paragraph(text)
            if save_immediately and self.file_path:
                self.save_word()
            elif save_immediately and not self.file_path:
                raise Exception("未指定保存路径，无法立即保存")
            return True
        except Exception as e:
            raise Exception(f"添加段落失败：{str(e)}")

    def read_paragraphs(self) -> List[str]:
        """读取所有段落文本"""
        if not self.document:
            raise Exception("请先新建或打开Word文档")
        try:
            paragraphs = [para.text for para in self.document.paragraphs if para.text.strip()]
            return paragraphs
        except Exception as e:
            raise Exception(f"读取段落失败：{str(e)}")

    def add_table(self, rows: int, cols: int, data: Optional[List[List[Any]]] = None) -> bool:
        """添加表格并填充数据（可选）"""
        if not self.document:
            raise Exception("请先新建或打开Word文档")
        try:
            table = self.document.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'  # 设置表格样式（带边框）
            # 填充数据
            if data and len(data) <= rows and all(len(row) <= cols for row in data):
                for i, row_data in enumerate(data):
                    table_row = table.rows[i]
                    for j, cell_data in enumerate(row_data):
                        table_row.cells[j].text = str(cell_data)
            return True
        except Exception as e:
            raise Exception(f"添加表格失败：{str(e)}")

    def save_word(self, file_path: Optional[str] = None) -> bool:
        """保存Word文档"""
        if not self.document:
            raise Exception("请先新建或打开Word文档")
        try:
            target_path = file_path or self.file_path
            if not target_path:
                raise Exception("请指定Word保存路径")
            self.document.save(target_path)
            self.file_path = target_path
            return True
        except Exception as e:
            raise Exception(f"保存Word失败：{str(e)}")

    def close_word(self) -> None:
        """关闭Word文档，释放资源"""
        if self.document:
            self.document = None
            self.file_path = None